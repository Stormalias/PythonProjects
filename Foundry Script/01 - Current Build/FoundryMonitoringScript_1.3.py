'''
Foundry Report Automation Script by Praval Visvanath 

This script ingests data from 2 csv files, checking for errors in the MTD and YTD for different benches, and prints a detailed
table of the errors in the portfolios of each bench. This output is stored in a word document, the results of which can be directly
pasted into the Teams chat to report to CR.

##################################################################################################################################################
INITIAL SETUP AND HOW TO USE:
    
    - The script requires some additional libraries to be installed, namely:
        > docx
        > ODBC
        > pandas
        
    This can be done via pip and the requirements.txt file, thusly:
        - Open command prompt
        - Navigate to the folder in which this Foundry script is located. Then run pip install using the requirements.txt
            > cd $DIRECTORY
            > pip install -r requirements.txt
    
    This will install all the packages listed in the text file. Depending on when this is done it may uninstall and reinstall packages according
    to the versions
    
    - As of 1.3 requires ODBC driver and a token to be made. Links for the one time setup will be provided separately (or ask the author)
    
    - The script directory structure should already be in place when downloaded. However to verify you should have the following structure:
    
    1.3\
        > Afternoon\
            > 
        > Morning\
            > 
        > __init__.py
        > csvExtractor.py
        > FoundryMonitoringScript_1.3.py
        > requirements.txt
    
    - If the above structure does not align to what you see, change it so it does. The Output files will appear in the 1.3\ directory
    
    - To use the script once setup, run FoundryMonitoringScript_1.3.py
    
    - You will be provided with 3 options. Pick the number according to what you are running and hit enter.
    
    - The script will run and once successful, will create the output Word document in 1.3\
    
    - Open the word document and enjoy your output
    
'''


'''
CHANGELOG
    Version 1.3:
        - Added csvExtractor package, which queries the Foundry DB using token and generates local csv. Script then continues as per normal
        - Requires user to install ODBC driver, and create DSN. Details in above section and knowledge base
        - Renamed output files so that they appear in order, to prevent it looking messy in the directory
        - Changed the introduction section to be more informative as a setup guide
        - Added commentary for csvExtractor.py
        - General clean up of code
    Version 1.2:
        - Changed the need for having 2 csv files. Instead users will now export all data from the pivot table, instead of just selected
        errors.
    Version 1.1:
        - Created functions to read CSVs from main export and contour export with date for selected errors, and then create a word doc with
        appropriate headers and data to paste into teams chat for CR
        - Added edge case for no errors to report
        - Three types of reports, as required, with monetary data in tables formatted as per request for CR with commas for every thousand places
        - Menu when running script, to choose which report runs

    Version 1.0
        - Created basic function to read from CSV of main export, and then create lines that read each bench with an error, followed by the 
        MTD and YTD error
'''

##################################################################################################################################################
#Start of script
##################################################################################################################################################


#Import Packages

import csvExtractor
import csv
import ctypes
import datetime
import docx
import pandas as pd
import sys


#Variable Declarations
mainFolder = sys.path[0]


#Function to find previous business day
def get_previous_business_day():
    today = pd.Timestamp.today()
    previous_day = today - pd.DateOffset(days=1)
    previous_business_day = pd.bdate_range(end=previous_day, periods=1)[0]
    return previous_business_day.strftime("%Y-%m-%d")


#Morning Report function
def mornReport(code = 0):
    if code == 1: #Afternoon Report assumes the morning report has been done
        pass
    else:
        csvExtractor.MorningCsv()
        
    #Opening file
    with open(mainFolder + "/Morning/COB_" + str(get_previous_business_day()) + ".csv") as file:
        #Reading Data
        Data = list(csv.reader(file))

        #Creating word document, with exception for singular afternoon report
        mydoc = docx.Document()
        if code == 2 :
            mydoc.add_paragraph("Hi Team,\n\nAll files have been processed.")
        elif code == 0 :
            mydoc.add_paragraph("Hi Team,\n\nTitan and Endur files have been processed.")
        
        #Dictionary mapping bench with errors to MTD and YTD error, to compare for afternoonReport, as well as to generate sums
        benchErrorDict = {}
        
        #Find errors, and add values to dictionary, appending to final sum at the end of loop
        for row in Data:
        
            if row[13] == 'ERROR':
                mtd = float(row[9])
                ytd = float(row[10])
                bench = row[1]
                
                if bench not in benchErrorDict:
                    benchErrorDict[bench] = [mtd, ytd]
                else:
                    benchErrorDict[bench][0] += mtd
                    benchErrorDict[bench][1] += ytd
            
        #Printing final sums
        for (bench, [mtdSum, ytdSum]) in benchErrorDict.items():
            mydoc.add_paragraph(f"\n\nErrors found in {bench} with MTD of {mtdSum} & YTD of {ytdSum} from the following portfolios:")
                    
            #Create table
            table = mydoc.add_table(1,10)
            line = table.rows[0].cells
            headers = Data[0][1:11]
            
            #Populating headers
            n = 0
            while n < 10:
                line[n].text = headers[n]
                n += 1
            
            #Print table with individual values
            for entry in Data:
                if str(entry[1]) == str(bench) and str(entry[13]) == 'ERROR':
                    line = table.add_row().cells
                    
                    i = 0
                    while i < 10:
                        if i > 2: #Skips over checking for negative values on columns with string outputs (non monetary)
                            val = '%.2f' % float(entry[i + 1])
                            line[i].text = str(format(float(val), ","))

                        else:
                            line[i].text = str(entry[i + 1])
                        i += 1
        
        if benchErrorDict == {}: #Edge case where there are no errors
            mydoc.add_paragraph("No errors to report this morning.")
            
        #Save output file or return dictionary for afternoonReport
        if code == 2 :
            mydoc.save(mainFolder + "/03 - Singular Foundry Report.docx")
        elif code == 0 :
            mydoc.save(mainFolder + "/01 - Morning Foundry Report.docx")
        elif code == 1 :
            return benchErrorDict


def afternoonReport():
    csvExtractor.AfternoonCsv()
    morningDict = mornReport(1) #Runs morning report to recreate Dictionary to compare
    #Opening file
    with open(mainFolder + "/Afternoon/COB_" + str(get_previous_business_day()) + ".csv") as file:
        #Reading Data
        Data = list(csv.reader(file))
        
        #Creating word document
        mydoc = docx.Document()
        mydoc.add_paragraph("Hi Team,\n\nAll files have been processed.")
        
        #Empty dictionary and lists for later use
        benchErrorDict = {}
        clearedErrors = []
        unchangedErrors = []
        
        #Find errors, and add values to dictionary, appending to final sum at the end of loop
        for row in Data:
            if row[13] == 'ERROR':
                mtd = float(row[9])
                ytd = float(row[10])
                bench = row[1]
                
                if bench not in benchErrorDict:
                    benchErrorDict[bench] = [mtd, ytd]
                else:
                    benchErrorDict[bench][0] += mtd
                    benchErrorDict[bench][1] += ytd
        
        
        #Checks for errors being cleared
        for key in morningDict.keys():
            if not key in benchErrorDict:
                clearedErrors.append(key)
        
        
        #Checks for no change in error
        for item in morningDict.items():
            if item in benchErrorDict.items():
                unchangedErrors.append(item[0])
        
        
        #Prints respective lines for unchanged error and cleared errors
        if unchangedErrors:
            mydoc.add_paragraph(f"No change in {', '.join(unchangedErrors)} numbers.")
        if clearedErrors:
            mydoc.add_paragraph(f"Errors with {', '.join(clearedErrors)} have been cleared.")
        
                
        #Checks for update in errors
        for (bench, [mtdSum, ytdSum]) in morningDict.items():
            
            if bench not in clearedErrors and bench not in unchangedErrors:
                #Heading over table, describing overall errors for the bench
                mydoc.add_paragraph(f"\n\nErrors found in {bench} with MTD of {benchErrorDict[bench][0]} & YTD of {benchErrorDict[bench][1]} from the following portfolios:")
                
                
                #Create table
                table = mydoc.add_table(1,10)
                line = table.rows[0].cells
                headers = Data[0][1:11]
        
                #Populating headers
                n = 0
                while n < 10:
                    line[n].text = headers[n]
                    n += 1
                
                #Print table with individual values
                for entry in Data:
                    if str(entry[1]) == str(bench) and str(entry[13]) == 'ERROR': #Create new row per match
                        line = table.add_row().cells
                        
                        i = 0
                        while i < 10:
                            if i > 2: #Monetary columns
                                val = '%.2f' % float(entry[i + 1])
                                line[i].text = str(format(float(val), ","))
                                
                            else: #Non-monetary columns 
                                line[i].text = str(entry[i + 1])
                            i += 1
        
        if benchErrorDict == {} and morningDict == {}: #Edge case where there are no errors
            mydoc.add_paragraph("No errors to report today.")
            
        #Save output file
        mydoc.save(mainFolder+ "/02 - Afternoon Foundry Report.docx")


def Menu(): #Menu to select report type
    reportNum = input("Is this the Morning(1), Afternoon(2) report or a Singular(3) report?\ninput 1,2 or 3 and hit Enter: ")
    if reportNum == '1' :
        mornReport()
        ctypes.windll.user32.MessageBoxW(0,"Morning Foundry Report Generated", "Success", 0)
    elif reportNum == '2' :
        afternoonReport()
        ctypes.windll.user32.MessageBoxW(0,"Afternoon Foundry Report Generated", "Success", 0)
    elif reportNum == '3' :
        mornReport(2)
        ctypes.windll.user32.MessageBoxW(0,"Singular Foundry Report Generated", "Success", 0)
    else :
        print("Incorrect entry. Enter 1 for Morning report, 2 for Afternoon report and 3 for Singular report\nOr hit control + C to quit\n\n")
        Menu()
        

if __name__ == "__main__":
    Menu() #Initialise
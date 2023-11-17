'''
Foundry Report Automation Script by Praval Visvanath 

This script ingests data from 2 csv files, checking for errors in the MTD and YTD for different benches, and prints a detailed
table of the errors in the portfolios of each bench. This output is stored in a word document, the results of which can be directly
pasted into the Teams chat to report to CR.

Things to note:
    - The script requires some libraries to be installed, namely:
        > docx
    This can be done via command prompt and pip3, with the commands:
        > pip3 install python-docx
    respectively.
    
    - The contour report file must be for today's date 
    
    - The script must be in the same folder as which the input and output files are to be expected.
'''


'''
CHANGELOG
    Version 1.2:
        - Changed the need for having 2 csv files. Instead users will now export all data from the pivot table, instead of just selected
        errors.
    Version 1.1:
        - Created functions to read CSVs from main export and contour export with date for selected errors, and then create a word doc with
        appropriate headers and data to paste into teams chat for CR
        - Added edge case for no errors to report
        - Three types of reports, as required, with monetary data in tables formatted as per request for CR with commas for every thousand places
        - Menu when running script, to choose which report runs

    Version 1.0
        - Created basic function to read from CSV of main export, and then create lines that read each bench with an error, followed by the 
        MTD and YTD error
'''


#Start of script
import csv
import docx
import datetime
import ctypes
import sys

#Variable Declarations
mainFolder = sys.path[0] 
td = datetime.datetime.now()
todayDate = td.strftime("%m-%d-%Y")


def mornReport(code = 0):
    #Opening file
    with open(mainFolder + "/Morning/contour-export-" + todayDate + ".csv") as file:
        #Reading Data
        Data = list(csv.reader(file))

        #Creating word document, with exception for singular afternoon report
        mydoc = docx.Document()
        if code == 2 :
            mydoc.add_paragraph("Hi Team,\n\nAll files have been processed.")
        elif code == 0 :
            mydoc.add_paragraph("Hi Team,\n\nTitan and Endur files have been processed.")
        
        #Dictionary mapping bench with errors to MTD and YTD error, to compare for afternoonReport, as well as to generate sums
        benchErrorDict = {}
        
        #Find errors, and add values to dictionary, appending to final sum at the end of loop
        for row in Data:
        
            if row[13] == 'ERROR':
                mtd = float(row[9])
                ytd = float(row[10])
                bench = row[1]
                
                if bench not in benchErrorDict:
                    benchErrorDict[bench] = [mtd, ytd]
                else:
                    benchErrorDict[bench][0] += mtd
                    benchErrorDict[bench][1] += ytd
            
        #Printing final sums
        for (bench, [mtdSum, ytdSum]) in benchErrorDict.items():
            mydoc.add_paragraph(f"\n\nErrors found in {bench} with MTD of {mtdSum} & YTD of {ytdSum} from the following portfolios:")
                    
            #Create table
            table = mydoc.add_table(1,10)
            line = table.rows[0].cells
            headers = Data[0][1:11]
            
            #Populating headers
            n = 0
            while n < 10:
                line[n].text = headers[n]
                n += 1
            
            #Print table with individual values
            for entry in Data:
                if str(entry[1]) == str(bench) and str(entry[13]) == 'ERROR':
                    line = table.add_row().cells
                    
                    i = 0
                    while i < 10:
                        if i > 2: #Skips over checking for negative values on columns with string outputs (non monetary)
                            val = '%.2f' % float(entry[i + 1])
                            line[i].text = str(format(float(val), ","))

                        else:
                            line[i].text = str(entry[i + 1])
                        i += 1
        
        if benchErrorDict == {}: #Edge case where there are no errors
            mydoc.add_paragraph("No errors to report this morning.")
            
        #Save output file or return dictionary for afternoonReport
        if code == 2 :
            mydoc.save(mainFolder + "/Singular Foundry Report.docx")
        elif code == 0 :
            mydoc.save(mainFolder + "/Morning Foundry Report.docx")
        elif code == 1 :
            return benchErrorDict


def afternoonReport():
    
    morningDict = mornReport(1)
    #Opening file
    with open(mainFolder + "/Afternoon/contour-export-" + todayDate + ".csv") as file:
        #Reading Data
        Data = list(csv.reader(file))
        
        #Creating word document
        mydoc = docx.Document()
        mydoc.add_paragraph("Hi Team,\n\nAll files have been processed.")
        
        #Empty dictionary and lists for later use
        benchErrorDict = {}
        clearedErrors = []
        unchangedErrors = []
        
        #Find errors, and add values to dictionary, appending to final sum at the end of loop
        for row in Data:
            if row[13] == 'ERROR':
                mtd = float(row[9])
                ytd = float(row[10])
                bench = row[1]
                
                if bench not in benchErrorDict:
                    benchErrorDict[bench] = [mtd, ytd]
                else:
                    benchErrorDict[bench][0] += mtd
                    benchErrorDict[bench][1] += ytd
        
        
        #Checks for errors being cleared
        for key in morningDict.keys():
            if not key in benchErrorDict:
                clearedErrors.append(key)
        
        
        #Checks for no change in error
        for item in morningDict.items():
            if item in benchErrorDict.items():
                unchangedErrors.append(item[0])
        
        
        #Prints respective lines for unchanged error and cleared errors
        if unchangedErrors:
            mydoc.add_paragraph(f"No change in {', '.join(unchangedErrors)} numbers.")
        if clearedErrors:
            mydoc.add_paragraph(f"Errors with {', '.join(clearedErrors)} have been cleared.")
        
                
        #Checks for update in errors
        for (bench, [mtdSum, ytdSum]) in morningDict.items():
            
            if bench not in clearedErrors and bench not in unchangedErrors:
                #Heading over table, describing overall errors for the bench
                mydoc.add_paragraph(f"\n\nErrors found in {bench} with MTD of {benchErrorDict[bench][0]} & YTD of {benchErrorDict[bench][1]} from the following portfolios:")
                
                
                #Create table
                table = mydoc.add_table(1,10)
                line = table.rows[0].cells
                headers = Data[0][1:11]
        
                #Populating headers
                n = 0
                while n < 10:
                    line[n].text = headers[n]
                    n += 1
                
                #Print table with individual values
                for entry in Data:
                    if str(entry[1]) == str(bench) and str(entry[13]) == 'ERROR': #Create new row per match
                        line = table.add_row().cells
                        
                        i = 0
                        while i < 10:
                            if i > 2: #Monetary columns
                                val = '%.2f' % float(entry[i + 1])
                                line[i].text = str(format(float(val), ","))
                                
                            else: #Non-monetary columns 
                                line[i].text = str(entry[i + 1])
                            i += 1
        
        if benchErrorDict == {} and morningDict == {}: #Edge case where there are no errors
            mydoc.add_paragraph("No errors to report today.")
            
        #Save output file
        mydoc.save(mainFolder+ "/Afternoon Foundry Report.docx")
        

def Menu(): #Menu to select report type
    reportNum = input("Is this the Morning(1), Afternoon(2) report or a Singular(3) report?\ninput 1,2 or 3 and hit Enter: ")
    if reportNum == '1' :
        mornReport()
        ctypes.windll.user32.MessageBoxW(0,"Morning Foundry Report Generated", "Success", 0)
    elif reportNum == '2' :
        afternoonReport()
        ctypes.windll.user32.MessageBoxW(0,"Afternoon Foundry Report Generated", "Success", 0)
    elif reportNum == '3' :
        mornReport(2)
        ctypes.windll.user32.MessageBoxW(0,"Singular Foundry Report Generated", "Success", 0)
    else :
        print("Incorrect entry. Enter 1 for Morning report, 2 for Afternoon report and 3 for Singular report\nOr hit control + C to quit\n\n")
        Menu()
        
Menu() #Initialise
'''
Foundry Report Automation Script by Praval Visvanath 

This script ingests data from 2 csv files, checking for errors in the MTD and YTD for different benches, and prints a detailed
table of the errors in the portfolios of each bench. This output is stored in a word document, the results of which can be directly
pasted into the Teams chat to report to CR.

Things to note:
    - The script requires some libraries to be installed, namely:
        > docx
    This can be done via command prompt and pip3, with the commands:
        > pip3 install python-docx
    respectively.
    
    - The contour report file must be for today's date 
    
    - The script must be in the same folder as which the input and output files are to be expected.
'''


'''
CHANGELOG
    Version 1.1:
        - Created functions to read CSVs from main export and contour export with date for selected errors, and then create a word doc with
        appropriate headers and data to paste into teams chat for CR
        - Added edge case for no errors to report
        - Three types of reports, as required, with monetary data in tables formatted as per request for CR with commas for every thousand places
        - Menu when running script, to choose which report runs

    Version 1.0
        - Created basic function to read from CSV of main export, and then create lines that read each bench with an error, followed by the 
        MTD and YTD error
'''


#Start of script
import csv
import docx
import datetime
import ctypes
import sys

#Variable Declarations
mainFolder = sys.path[0] 
td = datetime.datetime.now()
todayDate = td.strftime("%m-%d-%Y")


def mornReport(code = 0):
    #Opening files 
    with open(mainFolder + "/Morning/export.csv") as errorFile, open(mainFolder + "/Morning/contour-export-" + todayDate + ".csv") as contour:
        #Reading Data
        portfolioData = list(csv.reader(contour))
        errorCheck = list(csv.reader(errorFile))

        #Creating word document, with exception for singular afternoon report
        mydoc = docx.Document()
        if code == 2 :
            mydoc.add_paragraph("Hi Amanda / Shing Yee / Qiu Ting,\n\nAll files have been processed.")
        elif code == 0 :
            mydoc.add_paragraph("Hi Amanda / Shing Yee / Qiu Ting,\n\nTitan and Endur files have been processed.")
        
        #Dictionary mapping bench with errors to MTD and YTD error, to compare for afternoonReport
        benchErrorDict = {}
        
        #First loop to find errors
        errorCount = 0
        for row in errorCheck:
            if row[1] == 'ERROR':
                bench = row[0]
                mtdSum = row[-2]
                ytdSum = row[-1]
                benchErrorDict[bench] = [mtdSum, ytdSum]
                
                #Heading over table, describing overall errors for the bench
                mydoc.add_paragraph(f"\n\nErrors found in {bench} with MTD of {mtdSum} & YTD of {ytdSum} from the following portfolios:")
                
                #Create table
                table = mydoc.add_table(1,10)
                line = table.rows[0].cells
                headers = portfolioData[0][1:11]
                
                #Populating headers
                n = 0
                while n < 10:
                    line[n].text = headers[n]
                    n += 1
                
                #Second loop, to print portfolios of bench with error
                for entry in portfolioData:
                    if str(entry[1]) == str(bench):
                        line = table.add_row().cells
                        
                        i = 0
                        while i < 10:
                            if i > 2: #Skips over checking for negative values on columns with string outputs (non monetary)
                                val = '%.2f' % float(entry[i + 1])
                                line[i].text = str(format(float(val), ","))

                            else:
                                line[i].text = str(entry[i + 1])
                            i += 1
        
        if benchErrorDict == {}: #Edge case where there are no errors
            mydoc.add_paragraph("No errors to report this morning.")
            
        #Save output file or return dictionary for afternoonReport
        if code == 2 :
            mydoc.save(mainFolder + "/Singular Foundry Report.docx")
        elif code == 0 :
            mydoc.save(mainFolder + "/Morning Foundry Report.docx")
        elif code == 1 :
            return benchErrorDict


def afternoonReport():
    
    morningDict = mornReport(1)
    #Opening files 
    with open(mainFolder + "/Afternoon/export.csv") as errorFile, open(mainFolder + "/Afternoon/contour-export-" + todayDate + ".csv") as contour:
        #Reading Data
        portfolioData = list(csv.reader(contour))
        errorCheck = list(csv.reader(errorFile))
        
        #Creating word document
        mydoc = docx.Document()
        mydoc.add_paragraph("Hi Amanda / Shing Yee / Qiu Ting,\n\nAll files have been processed.")
        
        benchErrorDict = {}
        clearedErrors = []
        unchangedErrors = []
        
        #First loop to find errors
        for row in errorCheck:
            if row[1] == 'ERROR':
                bench = row[0]
                mtdSum = row[-2]
                ytdSum = row[-1]
                benchErrorDict[bench] = [mtdSum, ytdSum]
        
        #Checks for errors being cleared
        for key in morningDict.keys():
            if not key in benchErrorDict:
                clearedErrors.append(key)
        
        
        #Checks for no change in error
        for item in morningDict.items():
            if item in benchErrorDict.items():
                unchangedErrors.append(item[0])
        
        
        #Prints respective lines for unchanged error and cleared errors
        if unchangedErrors:
            mydoc.add_paragraph(f"No change in {', '.join(unchangedErrors)} numbers.")
        if clearedErrors:
            mydoc.add_paragraph(f"Errors with {', '.join(clearedErrors)} have been cleared.")
                
        #Checks for update in errors
        for (bench, [mtdSum, ytdSum]) in morningDict.items():
            if bench not in clearedErrors and bench not in unchangedErrors:
                #Heading over table, describing overall errors for the bench
                mydoc.add_paragraph(f"\n\nErrors found in {bench} with MTD of {benchErrorDict[bench][0]} & YTD of {benchErrorDict[bench][1]} from the following portfolios:")
        
                #Create table
                table = mydoc.add_table(1,10)
                line = table.rows[0].cells
                headers = portfolioData[0][1:11]
        
                #Populating headers
                n = 0
                while n < 10:
                    line[n].text = headers[n]
                    n += 1
                
                #Second loop, to print portfolios of bench with error
                for entry in portfolioData:
                    if str(entry[1]) == str(bench): #Create new row per match
                        line = table.add_row().cells
                        
                        i = 0
                        while i < 10:
                            if i > 2: #Monetary columns
                                val = '%.2f' % float(entry[i + 1])
                                line[i].text = str(format(float(val), ","))
                                
                            else: #Non-monetary columns 
                                line[i].text = str(entry[i + 1])
                            i += 1

        if benchErrorDict == {} and morningDict == {}: #Edge case where there are no errors
            mydoc.add_paragraph("No errors to report today.")
            
        #Save output file
        mydoc.save(mainFolder+ "/Afternoon Foundry Report.docx")

def Menu(): #Menu to select report type
    reportNum = input("Is this the Morning(1), Afternoon(2) report or a Singular(3) report?\ninput 1,2 or 3 and hit Enter: ")
    if reportNum == '1' :
        mornReport()
        ctypes.windll.user32.MessageBoxW(0,"Morning Foundry Report Generated", "Success", 0)
    elif reportNum == '2' :
        afternoonReport()
        ctypes.windll.user32.MessageBoxW(0,"Afternoon Foundry Report Generated", "Success", 0)
    elif reportNum == '3' :
        mornReport(2)
        ctypes.windll.user32.MessageBoxW(0,"Singular Foundry Report Generated", "Success", 0)
    else :
        print("Incorrect entry. Enter 1 for Morning report, 2 for Afternoon report and 3 for Singular report\nOr hit control + C to quit\n\n")
        Menu()
        
Menu() #Initialise
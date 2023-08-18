'''
Foundry Report Automation Script by Praval Visvanath 

This script ingests data from 2 csv files, checking for errors in the MTD and YTD for different benches, and prints a detailed
table of the errors in the portfolios of each bench. This output is stored in a word document, the results of which can be directly
pasted into the Teams chat to report to CR.

Things to note:
    - The script requires some libraries to be installed, namely:
        > csv
        > docx
    This can be done via command prompt and pip3, with the commands:
        > pip3 install csv
        > pip3 install python-docx
    respectively.
    
    - The contour report file must be for today's date 
    
    - The script must be in the same folder as which the input and output files are to be expected.
'''

#Start of script
import csv
import docx
from docx.shared import RGBColor
import datetime
import os
import ctypes
import sys

#Variable Declarations
mainFolder = sys.path[0]
td = datetime.datetime.now()
todayDate = td.strftime("%m-%d-%Y")
print(mainFolder)
print(todayDate)

def morningReport():
    #Opening files 
    with open(mainFolder + "/export.csv") as errorFile, open(mainFolder + "/contour-export-" + todayDate + ".csv") as contour:
        #Reading Data
        portfolioData = list(csv.reader(contour))
        errorCheck = list(csv.reader(errorFile))

        #Creating word document
        mydoc = docx.Document()
        mydoc.add_paragraph("Hi Team,\n\nPlatform files have been processed.")

        #First loop to find errors
        for row in errorCheck:
            print("Entering top loop") #debug line
            if row[1] == 'ERROR':
                bench = row[0]
                mtdSum = row[-2]
                ytdSum = row[-1]
                print(f"Top loop if check found error in {bench}") #debug line
                mydoc.add_paragraph(f"\n\nErrors found in {bench} with MTD of {mtdSum} & YTD of {ytdSum} from the following portfolios:")
                
                table = mydoc.add_table(1,10)
                line = table.rows[0].cells
                
                headers = portfolioData[0][1:11]
                
                n = 0
                #Populating headers
                while n < 10:
                    line[n].text = headers[n]
                    n += 1
                
                #Second loop, to print portfolios of bench with error
                for entry in portfolioData:
                    print(f"Inner loop looking for app bench. Currently matching {entry[1]} and {bench}") #debug line
                    if str(entry[1]) == str(bench):
                        print("Found match") #debug line
                        print(f"This current bench is {bench}") #debug line
                        line = table.add_row().cells
                        
                        i = 0
                        while i < 10:
                            if i > 2: #Skips over checking for negative values on columns with string outputs (non monetary)
                                if float(entry[i + 1]) < 0:
                                    
                                    val = '%.2f' % float(entry[i + 1])
                                    line[i].text = str(format(float(val), ","))
                                            
                                else :
                                    val = '%.2f' % float(entry[i + 1])
                                    line[i].text = str(format(float(val), ","))
                                    
                            
                            else :
                                line[i].text = str(entry[i + 1])
                            i += 1

        #Save output file
        mydoc.save(mainFolder+ "/Foundry Report.docx")
        print("Success") #debug line
        ctypes.windll.user32.MessageBoxW(0,"Foundry Report Generated", "Success", 0)




reportNum = input("Is this 1st or 2nd report: ")
print(reportNum)
print(type(reportNum))

if reportNum == '1' :
    morningReport()
else :
    ctypes.windll.user32.MessageBoxW(0,"2nd report", "haven do", 0)